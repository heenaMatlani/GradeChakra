from docx import Document
from docx.shared import Pt
from decimal import Decimal
import datetime

def set_font_size(cell, font_size=8):
    """
    Set the font size for all text within a cell to the specified font size in points.
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)  # Set font size explicitly in points

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(0)

def replace_placeholders_in_table(table, data_dict, font_size=8):
    """
    Replace placeholders in a single table using data from data_dict and set font size.
    """
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text
            for key, value in data_dict.items():
                placeholder = f"<<{key}>>"
                if placeholder in cell_text:
                    cell.text = cell_text.replace(placeholder, str(value) if value else '')

            # Set the font size for each cell after replacement
            set_font_size(cell, font_size)

    # Clear any remaining placeholders in the table and set font size
    for row in table.rows:
        for cell in row.cells:
            for key in data_dict.keys():
                placeholder = f"<<{key}>>"
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, '')
            set_font_size(cell, font_size)

def replace_course_data_in_table(table, courses, font_size=8):
    """
    Fill each row in the table with course information from courses list and set font size.
    """
    for course in courses:
        row_filled = False
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                for key, value in course.items():
                    placeholder = f"<<{key}>>"
                    if placeholder in cell_text:
                        cell.text = cell_text.replace(placeholder, str(value) if value else '')
                        row_filled = True
            if row_filled:
                break  # Move to the next course

    # Clear any remaining course placeholders and set font size
    for row in table.rows:
        for cell in row.cells:
            for key in courses[0].keys():
                placeholder = f"<<{key}>>"
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, '')
            set_font_size(cell, font_size)

def replace_placeholders(doc, data, font_size=8):
    """
    Replace placeholders in both paragraphs and tables in the Word document with actual data.
    """
    # Static data replacement in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"<<{key}>>"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value) if value else '')

    # Process each table for placeholders
    for table in doc.tables:
        replace_placeholders_in_table(table, data, font_size)

        # Process each semester's courses
        for semester_data in data.get('semesters', []):
            replace_course_data_in_table(table, semester_data.get('courses', []), font_size)
            semester_metadata = {k: semester_data[k] for k in ['Semester_Start', 'Acad_Start', 'Semester_End', 'Acad_End']}
            replace_placeholders_in_table(table, semester_metadata, font_size)

def generate_docx_report(template_path, data, output_docx_path, font_size=8):
    """
    Generate a .docx report from a template and set font size.
    """
    doc = Document(template_path)
    replace_placeholders(doc, data, font_size)
    doc.save(output_docx_path)

# Sample data for placeholder replacement

# Function to recursively convert all values to strings
def convert_to_string(data):
    if isinstance(data, dict):
        return {key: convert_to_string(value) for key, value in data.items()}
    elif isinstance(data, list):
        return [convert_to_string(item) for item in data]
    elif isinstance(data, (Decimal, int, float, datetime.date)):
        return str(data)
    else:
        return data

data = {
    'Name': 'Heena Matlani',
    'Program_Name': 'Bachelor of Technology',
    'Roll_Number': 'B202001001',
    'Batch_Year': 2020,
    'Department_Name': 'Computer Science and Engineering',
    'Total_Semesters': 8,
    'Date': datetime.date(2024, 11, 14),
    'semesters': [
        {
            'Semester_Start': 'August',
            'Acad_Start': 2020,
            'Semester_End': 'December',
            'Acad_End': 2021,
            'courses': [
                {'Course_Name': 'Computer Programming', 'Course_Code': 'CS666', 'Course_Credit': 6, 'Course_Grade': 'AA'},
                {'Course_Name': 'English', 'Course_Code': 'HS101', 'Course_Credit': 4, 'Course_Grade': 'AB'},
                {'Course_Name': 'Mathematics I', 'Course_Code': 'MA101', 'Course_Credit': 8, 'Course_Grade': 'BC'},
                {'Course_Name': 'Electrical Circuit Analysis', 'Course_Code': 'EC102', 'Course_Credit': 8, 'Course_Grade': 'BB'},
                {'Course_Name': 'Digital Design Lab', 'Course_Code': 'EC110', 'Course_Credit': 3, 'Course_Grade': 'AA'},
                {'Course_Name': 'Digital Design', 'Course_Code': 'EC666', 'Course_Credit': 8, 'Course_Grade': 'AA'},
                {'Course_Name': 'Computer Programming Lab', 'Course_Code': 'CS110', 'Course_Credit': 5, 'Course_Grade': 'BB'},
                {'Course_Name': 'Computer Programming', 'Course_Code': 'CS666', 'Course_Credit': 6, 'Course_Grade': 'BC'}
            ]
        },
        {
            'Semester_Start': 'January',
            'Acad_Start': 2020,
            'Semester_End': 'May',
            'Acad_End': 2021,
            'courses': [
                {'Course_Name': 'Introduction to Politics', 'Course_Code': 'HS204', 'Course_Credit': 6, 'Course_Grade': 'AA'},
                {'Course_Name': 'Basic Electronic Lab', 'Course_Code': 'EC111', 'Course_Credit': 3, 'Course_Grade': 'AA'},
                {'Course_Name': 'Basic Electronic Circuits', 'Course_Code': 'EC103', 'Course_Credit': 8, 'Course_Grade': 'AA'},
                {'Course_Name': 'Computer Organization', 'Course_Code': 'CS104', 'Course_Credit': 8, 'Course_Grade': 'AB'},
                {'Course_Name': 'Data Structures Lab', 'Course_Code': 'CS111', 'Course_Credit': 3, 'Course_Grade': 'BB'},
                {'Course_Name': 'Data Structures', 'Course_Code': 'CS103', 'Course_Credit': 8, 'Course_Grade': 'BC'},
                {'Course_Name': 'Mathematics II', 'Course_Code': 'MA102', 'Course_Credit': 8, 'Course_Grade': 'AB'}
            ]
        }
    ],
    'SPI_Name': 'S.P.I',
    'CPI_Name': 'C.P.I',
    'SPI_CPI': {'Sem_1_spi': Decimal('8.71'), 'Sem_1_cpi': Decimal('8.71'), 'Sem_2_spi': Decimal('8.95'), 'Sem_2_cpi': Decimal('8.84')},
    'is_completed': 'INCOMPLETE'
}

# Convert all data to string
data = convert_to_string(data)
# Paths for the template and the output .docx
template_path = '/Users/heenamatlani/Downloads/GradeCard_Template_Final.docx'
output_docx_path = 'Generated_Grade_Report3.docx'

# Generate the .docx report
generate_docx_report(template_path, data, output_docx_path, font_size=8)
print("DOCX report generated successfully at:", output_docx_path)
